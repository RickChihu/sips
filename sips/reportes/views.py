import csv

from django.shortcuts import render
from django.http import HttpResponse
from io import BytesIO
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A2
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from reportlab.platypus.paragraph import Paragraph

from asuntos.models import Asunto

from cStringIO import StringIO
from docx import Document
from datetime import date, timedelta


def reports_view(request):
    return render(request, 'reportes/listado.html', {})


def export_asuntos_nuevos_csv(request):
    today = date.today()
    first_day_of_month = today.replace(day=1)

    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="users.csv"'

    writer = csv.writer(response)
    writer.writerow(['Asunto', 'Folio', 'Nombre', 'Curp', 'Domicilio', 'Fecha Creacion', 'Status', 'Agente Social'])

    asuntos = Asunto.objects.filter(added__gte=first_day_of_month)
    for asunto in asuntos:
        writer.writerow([
            asunto.asunto.name.encode('utf-8'),
            asunto.folio_().encode('utf-8'),
            asunto.nombre.encode('utf-8'),
            asunto.curp.encode('utf-8'),
            asunto.domicilio.encode('utf-8'),
            str(asunto.added).encode('utf-8'),
            asunto.status.encode('utf-8'),
            asunto.agente_social.username.encode('utf-8') if asunto.agente_social else '',
        ])

    return response


def export_asuntos_mes_pasado_csv(request):
    today = date.today()
    first_day_of_month = today.replace(day=1)
    first_day_of_month_before = today - timedelta(days=30)

    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="users.csv"'

    writer = csv.writer(response)
    writer.writerow(['Asunto', 'Folio', 'Nombre', 'Curp', 'Domicilio', 'Fecha Creacion', 'Status', 'Agente Social'])

    asuntos = Asunto.objects.filter(added__gte=first_day_of_month_before, added__lte=first_day_of_month)
    for asunto in asuntos:
        writer.writerow([
            asunto.asunto.name.encode('utf-8'),
            asunto.folio_().encode('utf-8'),
            asunto.nombre.encode('utf-8'),
            asunto.curp.encode('utf-8'),
            asunto.domicilio.encode('utf-8'),
            str(asunto.added).encode('utf-8'),
            asunto.status.encode('utf-8'),
            asunto.agente_social.username.encode('utf-8') if asunto.agente_social else '',
        ])

    return response


def export_asuntos_nuevos_word(request):
    document = Document()
    table = document.add_table(rows=1, cols=8)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Asunto'
    hdr_cells[1].text = 'Folio'
    hdr_cells[2].text = 'Nombre'
    hdr_cells[3].text = 'Curp'
    hdr_cells[4].text = 'Domicilio'
    hdr_cells[5].text = 'Fecha Creacion'
    hdr_cells[6].text = 'Status'
    hdr_cells[7].text = 'Agente Social'

    today = date.today()
    first_day_of_month = today.replace(day=1)

    asuntos = Asunto.objects.filter(added__gte=first_day_of_month)
    for asunto in asuntos:
        row_cells = table.add_row().cells
        row_cells[0].text = asunto.asunto.name
        row_cells[1].text = asunto.folio_()
        row_cells[2].text = asunto.nombre
        row_cells[3].text = asunto.nombre
        row_cells[4].text = asunto.nombre
        row_cells[5].text = asunto.nombre
        row_cells[6].text = asunto.nombre
        row_cells[7].text = asunto.nombre

    f = StringIO()
    document.save(f)
    length = f.tell()
    f.seek(0)

    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

    response['Content-Disposition'] = 'attachment; filename=example.docx'
    response['Content-Length'] = length
    return response


def export_asuntos_mes_pasado_word(request):
    document = Document()
    table = document.add_table(rows=1, cols=8)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Asunto'
    hdr_cells[1].text = 'Folio'
    hdr_cells[2].text = 'Nombre'
    hdr_cells[3].text = 'Curp'
    hdr_cells[4].text = 'Domicilio'
    hdr_cells[5].text = 'Fecha Creacion'
    hdr_cells[6].text = 'Status'
    hdr_cells[7].text = 'Agente Social'

    today = date.today()
    first_day_of_month = today.replace(day=1)
    first_day_of_month_before = today - timedelta(days=30)

    asuntos = Asunto.objects.filter(added__gte=first_day_of_month_before, added__lte=first_day_of_month)
    for asunto in asuntos:
        row_cells = table.add_row().cells
        row_cells[0].text = asunto.asunto.name
        row_cells[1].text = asunto.folio_()
        row_cells[2].text = asunto.nombre
        row_cells[3].text = asunto.nombre
        row_cells[4].text = asunto.nombre
        row_cells[5].text = asunto.nombre
        row_cells[6].text = asunto.nombre
        row_cells[7].text = asunto.nombre

    f = StringIO()
    document.save(f)
    length = f.tell()
    f.seek(0)

    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

    response['Content-Disposition'] = 'attachment; filename=example.docx'
    response['Content-Length'] = length
    return response


def export_asuntos_nuevos_pdf(request):
    response = HttpResponse(content_type='application/pdf')
    pdf_name = "reporte.pdf"  # llamado clientes
    buff = BytesIO()
    doc = SimpleDocTemplate(buff,
                            pagesize=A2,
                            rightMargin=40,
                            leftMargin=40,
                            topMargin=60,
                            bottomMargin=18,
                            )
    data = []
    styles = getSampleStyleSheet()
    header = Paragraph("Reporte de Asuntos nuevos en el mes", styles['Heading1'])
    data.append(header)
    headings = ('Asunto', 'Folio', 'Nombre', 'Curp', 'Domicilio', 'Status', 'Agente Social')
    today = date.today()
    first_day_of_month = today.replace(day=1)

    records = [(asunto.asunto.name, asunto.folio_(), asunto.nombre, asunto.curp, asunto.domicilio,
                asunto.status, asunto.agente_social.username if asunto.agente_social else '')
               for asunto in Asunto.objects.filter(added__gte=first_day_of_month)]

    t = Table([headings] + records)
    t.setStyle(TableStyle(
        [
            ('GRID', (0, 0), (6, -1), 1, colors.dodgerblue),
            ('LINEBELOW', (0, 0), (-1, 0), 2, colors.darkblue),
            ('BACKGROUND', (0, 0), (-1, 0), colors.dodgerblue)
        ]
    ))
    data.append(t)
    doc.build(data)
    response.write(buff.getvalue())
    buff.close()
    response['Content-Disposition'] = 'attachment; filename="reporte.pdf"'

    return response


def export_asuntos_mes_pasado_pdf(request):
    response = HttpResponse(content_type='application/pdf')
    pdf_name = "reporte.pdf"
    buff = BytesIO()
    doc = SimpleDocTemplate(buff,
                            pagesize=A2,
                            rightMargin=40,
                            leftMargin=40,
                            topMargin=60,
                            bottomMargin=18,
                            )
    data = []
    styles = getSampleStyleSheet()
    header = Paragraph("Reporte de Asuntos del mes pasado", styles['Heading1'])
    data.append(header)
    headings = ('Asunto', 'Folio', 'Nombre', 'Curp', 'Domicilio', 'Status', 'Agente Social')
    today = date.today()
    first_day_of_month = today.replace(day=1)
    first_day_of_month_before = today - timedelta(days=30)

    records = [(asunto.asunto.name, asunto.folio_(), asunto.nombre, asunto.curp, asunto.domicilio,
                asunto.status, asunto.agente_social.username if asunto.agente_social else '')
               for asunto in Asunto.objects.filter(added__gte=first_day_of_month_before, added__lte=first_day_of_month)]

    t = Table([headings] + records)
    t.setStyle(TableStyle(
        [
            ('GRID', (0, 0), (6, -1), 1, colors.dodgerblue),
            ('LINEBELOW', (0, 0), (-1, 0), 2, colors.darkblue),
            ('BACKGROUND', (0, 0), (-1, 0), colors.dodgerblue)
        ]
    ))
    data.append(t)
    doc.build(data)
    response.write(buff.getvalue())
    buff.close()
    response['Content-Disposition'] = 'attachment; filename="reporte.pdf"'

    return response
